"""Adapter `DocumentParser` para arquivos .docx (OOXML), via python-docx.

Produz o **mesmo** `ParsedDocument` do PDF, populando os atributos de que o
`DetectInjection` depende: cor da fonte (branco-no-branco), tamanho (fonte
minúscula), **texto oculto** (`w:vanish` → `in_hidden_ocg`) e metadados (core
properties). DOCX é fluxo de texto, sem coordenadas — então as checagens
**geométricas** (off-cropbox/OCG/ActualText) simplesmente não se aplicam. Cada
parágrafo vira um "bloco" (para o `visible_text` por bloco).

Cobertura v1: corpo, tabelas e cabeçalhos/rodapés. Caixas de texto (drawing XML)
e cores de tema ficam como limitação conhecida.
"""

from __future__ import annotations

from collections.abc import Iterator
from io import BytesIO

import docx
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from sherpi.contexts.document_integrity.domain.document import ParsedDocument, TextSpan
from sherpi.shared_kernel.errors import UntrustedDocumentError

_DEFAULT_SIZE_PT = 11.0  # tamanho assumido quando a run não declara (herdaria do estilo)
_BLACK = (0.0, 0.0, 0.0)
_MAX_PARAGRAPHS = 50_000  # guarda contra zip-bomb / documento abusivo


class DocxParser:
    """Implementa o port `DocumentParser` para .docx."""

    def parse(self, content: bytes, *, max_pages: int) -> ParsedDocument:
        try:
            document = docx.Document(BytesIO(content))
        except Exception as exc:  # arquivo hostil / corrompido
            raise UntrustedDocumentError(f"Falha ao abrir o DOCX: {exc}") from exc

        spans: list[TextSpan] = []
        for block, paragraph in enumerate(self._iter_paragraphs(document)):
            if block > _MAX_PARAGRAPHS:
                raise UntrustedDocumentError("DOCX excede o limite de parágrafos.")
            for run in paragraph.runs:
                span = self._run_to_span(run, block)
                if span is not None:
                    spans.append(span)

        return ParsedDocument(
            page_count=1,
            spans=spans,
            pages=[],  # sem geometria de página → off-cropbox/imagem não se aplicam
            metadata=self._extract_metadata(document),
            has_optional_content=False,
        )

    @staticmethod
    def _iter_paragraphs(document: DocxDocument) -> Iterator[Paragraph]:
        """Parágrafos do corpo + células de tabela + cabeçalhos/rodapés."""
        yield from document.paragraphs
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
        for section in document.sections:
            yield from section.header.paragraphs
            yield from section.footer.paragraphs

    @staticmethod
    def _run_to_span(run: Run, block: int) -> TextSpan | None:
        text = run.text
        if not text:
            return None
        font = run.font
        size = float(font.size.pt) if font.size is not None else _DEFAULT_SIZE_PT
        rgb = _BLACK
        if font.color is not None and font.color.rgb is not None:
            c = font.color.rgb  # RGBColor: tupla de inteiros 0..255
            rgb = (c[0] / 255, c[1] / 255, c[2] / 255)
        return TextSpan(
            text=text,
            rgb=rgb,
            size=size,
            bbox=(0.0, 0.0, 0.0, 0.0),
            page=0,
            block=block,
            in_hidden_ocg=bool(font.hidden),  # w:vanish — texto oculto
        )

    @staticmethod
    def _extract_metadata(document: DocxDocument) -> dict[str, str]:
        cp = document.core_properties
        fields = {
            "title": cp.title,
            "subject": cp.subject,
            "keywords": cp.keywords,
            "comments": cp.comments,
            "category": cp.category,
            "author": cp.author,
        }
        return {k: v for k, v in fields.items() if v}
