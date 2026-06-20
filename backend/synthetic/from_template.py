"""Gerador de petições sintéticas **a partir de templates reais** (.docx).

Diferente do `synthetic/builder.py` (que monta PDFs do zero, com texto pobre), este
módulo recebe um **template oficial** (`.docx` — ex.: os formulários de autoatendimento
do TJDFT em `docs/templates/`), preenche os campos com dados **sintéticos** (sem PII
real — decisão *synthetic-first* / LGPD) e produz um par `.docx` + `.pdf` com a petição
inicial já redigida, mais um *ground truth* no schema do `labels.json`.

Tipos de campo reconhecidos no template (mapeados em petições reais do DF):
- **placeholder** entre angle-brackets: ``<DIGITE O NOME DO AUTOR>``;
- **campo em branco**: ``nacionalidade:      ,`` (rótulo + 2+ espaços);
- **checkbox**: ``(  ) evadiu-se do local`` — marca uma opção por grupo;
- **bloco alternativo**: ``Situação 1 - ... / Situação 2 - ...`` — mantém um;
- **meta-instrução**: ``(sugestões de texto …)``, notas com ``*`` — removidas.

Uso:
    uv run python -m synthetic.from_template TEMPLATE.docx \\
        [--orientacao ORIENTACOES.pdf] [--n 3] [--seed 42] \\
        [--out data/synthetic] [--rito CIVEL] [--category clean] \\
        [--enrich-llm] [--keep-docx]

O preenchimento é **determinístico e semeado** por padrão (reprodutível, offline,
sem custo de token). ``--enrich-llm`` reescreve só a narrativa livre (fatos) via o
port `LLMProvider`, usando o `.pdf` de orientação como contexto.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import random
import re
import shutil
import subprocess
import tempfile
import unicodedata
from collections.abc import Callable, Iterator
from dataclasses import asdict, dataclass, field
from pathlib import Path

import docx
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from synthetic.entities import (
    Endereco,
    PessoaFisica,
    PessoaJuridica,
    make_distinct_pessoas,
    make_pessoa_juridica,
)

_BLACK = RGBColor(0x00, 0x00, 0x00)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# --- Padrões de slot -------------------------------------------------------------

# Placeholder: <...> SEM '<' ou '>' no miolo. Aceita parênteses internos
# (ex.: "<DIGITE O NOME DA CIDADE (FÓRUM)>").
_ANGLE = re.compile(r"<[^<>]+>")
# Espaço "de preenchimento": 2+ brancos de qualquer tipo, menos quebra de linha.
# Os templates do TJDFT usam EN SPACE (U+2002), não o espaço comum — por isso
# [^\S\n] (qualquer whitespace Unicode menos quebra de linha) em vez de [ ].
_WS = r"[^\S\n]"
# Campo em branco: "rótulo:" seguido de 2+ brancos. O rótulo não cruza vírgula
# (campos adjacentes "X:   , Y:   ," terminam em ", ") — isola o nome do campo.
_BLANK = re.compile(rf"([^:,\n]{{2,40}}?):{_WS}{{2,}}")
# Checkbox vazio: "(  )" (2+ brancos entre parênteses).
_CHECKBOX = re.compile(rf"\({_WS}{{2,}}\)")
# "h  min" do horário (ex.: por volta das   h  min).
_HORA = re.compile(rf"(?<=das){_WS}{{2,}}h{_WS}{{2,}}min")
# Blanks SEM rótulo: "R$   " (valor) e "nº   " (número de documento).
_RS_BLANK = re.compile(rf"R\$ ?{_WS}{{2,}}")
_NUM_BLANK = re.compile(rf"n[ºo°]\.?:?{_WS}{{2,}}")
# Placeholder que CRUZA fronteira de parágrafo: abre '<...' sem fechar / fecha '...>'.
_OPEN_ANGLE = re.compile(r"<[^<>]*$")
_CLOSE_ANGLE = re.compile(r"^[^<>]*>")

# Linhas de meta-instrução do template (não fazem parte da peça final).
_META_MARKERS = (
    "-autoatendimento-",
    "sugestões de texto",
    "adote um",
    "apague o outro",
    "recomenda-se",
    "marque mais de uma",
    "pode marcar mais de uma",
    "preencha",
    "instruç",
)


# --- Contexto de preenchimento (dados sintéticos coerentes por peça) -------------


def _money_brl(value: int) -> str:
    """Formata um inteiro como moeda BRL sem o prefixo 'R$' (ex.: 12.000,00)."""
    return f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


# Índices fixos das partes de uma peça (litisconsórcio até 2x2 + representante).
_AUTOR, _AUTOR_CONDUTOR, _REU, _REU_CONDUTOR, _TERCEIRO = range(5)


def _rua(e: Endereco) -> str:
    """Trecho de logradouro do endereço (a cidade/CEP vão em campos próprios)."""
    return f"{e.logradouro}, nº {e.numero}, {e.bairro}"


@dataclass
class _FillCtx:
    """Dados sintéticos de uma petição: partes (PF) e réu PJ vindos do catálogo
    `entities`, mais o estado da *parte corrente* — para que blocos de qualificação
    de partes diferentes não repitam nome/CPF/endereço — e campos coerentes da peça."""

    rng: random.Random
    personas: list[PessoaFisica] = field(default_factory=list)
    empresa_pj: PessoaJuridica | None = None
    cur: int = _AUTOR  # parte corrente; muda ao resolver o nome de cada parte
    use_empresa: bool = False  # True dentro do bloco do réu pessoa jurídica
    data: str = ""
    contrato: str = ""
    beneficio: str = ""
    agencia_conta: str = ""
    parcelas: int = 0
    veiculo_marca: str = ""
    veiculo_modelo: str = ""
    veiculo_cor: str = ""
    veiculo_ano: str = ""
    placa: str = ""
    valor_base: int = 0  # âncora monetária da peça (demais valores derivam dela)
    local: str = ""
    avarias: str = ""
    outras_info: str = ""
    fatos: str = ""  # narrativa (preenchida no modo --enrich-llm)

    def __post_init__(self) -> None:
        # Prenomes distintos entre as partes da MESMA peça (evita "Lucas ... Lucas ...").
        self.personas = make_distinct_pessoas(self.rng, 5)
        self.empresa_pj = make_pessoa_juridica(self.rng)
        self.data = f"{self.rng.randint(1, 28):02d}/{self.rng.randint(1, 12):02d}/2025"
        self.contrato = f"{self.rng.randint(100000, 999999)}"
        self.beneficio = f"{self.rng.randint(100, 999)}.{self.rng.randint(100, 999)}.{self.rng.randint(100, 999)}-{self.rng.randint(0, 9)}"
        self.agencia_conta = f"Banco 001, agência {self.rng.randint(1000, 9999)}, conta {self.rng.randint(10000, 99999)}-{self.rng.randint(0, 9)}"
        self.parcelas = self.rng.choice([12, 24, 36, 48, 60])
        self.veiculo_marca, self.veiculo_modelo = self.rng.choice(
            [
                ("Volkswagen", "Gol 1.6"),
                ("Chevrolet", "Onix"),
                ("Fiat", "Argo"),
                ("Hyundai", "HB20"),
            ]
        )
        self.veiculo_cor = self.rng.choice(["prata", "branco", "preto", "vermelho"])
        self.veiculo_ano = f"{self.rng.randint(2015, 2023)}/{self.rng.randint(2016, 2024)}"
        self.placa = f"{self.rng.choice('ABCDEFGH')}{self.rng.choice('ABCDEFGH')}{self.rng.choice('ABCDEFGH')}{self.rng.randint(0, 9)}{self.rng.choice('ABCDEFGH')}{self.rng.randint(10, 99)}"
        self.valor_base = self.rng.choice([3_500, 5_000, 8_000, 12_000, 18_000, 25_000])
        self.local = self.rng.choice(
            [
                "Av. das Nações, sentido Rodoviária",
                "EPTG, próximo ao viaduto",
                "Av. Hélio Prates, em frente ao terminal",
                "SIA, trecho 3",
            ]
        )
        self.avarias = self.rng.choice(
            [
                "traseira",
                "lateral dianteira esquerda",
                "para-choque dianteiro",
                "porta do motorista",
            ]
        )
        self.outras_info = (
            "o veículo da parte autora não possui seguro e o conserto ainda não foi realizado"
        )

    @property
    def current(self) -> PessoaFisica:
        return self.personas[self.cur]

    @property
    def empresa_entity(self) -> PessoaJuridica:
        assert self.empresa_pj is not None
        return self.empresa_pj

    @property
    def endereco(self) -> Endereco:
        """Endereço da parte ativa (empresa ré dentro do bloco do réu PJ)."""
        return self.empresa_entity.endereco if self.use_empresa else self.current.endereco

    @property
    def telefone(self) -> str:
        return self.empresa_entity.telefone if self.use_empresa else self.current.telefone

    @property
    def email(self) -> str:
        return self.empresa_entity.email if self.use_empresa else self.current.email

    @property
    def autor_nome(self) -> str:
        return self.personas[_AUTOR].nome

    @property
    def reu_nome(self) -> str:
        return self.personas[_REU].nome

    @property
    def autor_cpf(self) -> str:
        return self.personas[_AUTOR].cpf

    @property
    def terceiro_nome(self) -> str:
        return self.personas[_TERCEIRO].nome

    @property
    def reu_empresa(self) -> str:
        return self.empresa_entity.razao_social

    @property
    def reu_cnpj(self) -> str:
        return self.empresa_entity.cnpj

    def name_for(self, key: str) -> str:
        """Resolve um placeholder de NOME de parte → muda a parte corrente e devolve
        o nome dela. Distingue litisconsórcio por 'condutor' (2ª parte de cada lado)."""
        self.use_empresa = False
        condutor = "condutor" in key
        if "representante" in key or "preposto" in key:
            self.cur = _TERCEIRO
        elif re.search(r"reu|requerid|demandad|reclamad", key):
            self.cur = _REU_CONDUTOR if condutor else _REU
        else:  # lado autor (autor/requerente/proprietário/condutor/seu nome)
            self.cur = _AUTOR_CONDUTOR if condutor else _AUTOR
        return self.current.nome

    def empresa(self) -> str:
        """Réu pessoa jurídica: ativa o contexto da empresa (endereço/telefone/CNPJ
        seguintes saem DELA) e devolve a razão social."""
        self.use_empresa = True
        return self.empresa_entity.razao_social

    def money(self, key: str) -> str:
        """Valor monetário **estável por campo**: a mesma chave semântica devolve
        sempre o mesmo número; chaves distintas (parcela ≠ total ≠ dano moral)
        devolvem valores distintos, ancorados em `valor_base`."""
        bucket = int(hashlib.md5(key.encode("utf-8")).hexdigest(), 16) % 8
        factor = [0.1, 0.25, 0.5, 1.0, 1.0, 1.5, 2.0, 3.0][bucket]
        value = max(200, round(self.valor_base * factor / 50) * 50)
        return _money_brl(value)


# Registry: padrão sobre a **chave semântica** (placeholder/rótulo minúsculo, sem
# acento, sem "digite/digitar/aqui") -> gerador(ctx, key). A 1ª regra que casa vence.
_RULES: list[tuple[re.Pattern[str], Callable[[_FillCtx, str], str]]] = [
    # Nomes de parte primeiro: mudam a "parte corrente" da qual os dados pessoais
    # seguintes (CPF/RG/endereço) são extraídos — evita litisconsórcio repetido.
    (re.compile(r"representante|preposto"), lambda c, k: c.name_for(k)),
    # Réu PESSOA JURÍDICA: só com indicador EXPLÍCITO de PJ (razão social, fantasia,
    # banco, empresa). 'requerida/requerido' genérico NÃO entra aqui — pode ser
    # pessoa física (template "contra PESSOA FÍSICA") e cai na regra de nome.
    (
        re.compile(r"razao social|fantasia|\bbanco\b|empresa|pessoa juridica"),
        lambda c, k: c.empresa(),
    ),
    (re.compile(r"orgao publico|distrito federal|gdf"), lambda c, k: "Distrito Federal"),
    (re.compile(r"nome d[oa] acao|conhecimento"), lambda c, k: "conhecimento"),
    (
        re.compile(
            r"nome.*(autor|requerente|demandante|reclamante|reu|requerid|demandad|reclamad|proprietario|condutor)|seu nome|nome completo"
        ),
        lambda c, k: c.name_for(k),
    ),
    (re.compile(r"inscricao estadual"), lambda c, k: c.empresa_entity.inscricao_estadual),
    (re.compile(r"\bcnpj\b"), lambda c, k: c.reu_cnpj),
    (re.compile(r"\bcpf\b"), lambda c, k: c.current.cpf),
    (re.compile(r"\brg\b|identidade|carteira de ident"), lambda c, k: c.current.rg),
    (re.compile(r"expedidor|orgao exped|\buf\b|\bssp\b"), lambda c, k: c.current.orgao_expedidor),
    (re.compile(r"nascimento"), lambda c, k: c.current.data_nascimento),
    (re.compile(r"e-?mail"), lambda c, k: c.email),
    (re.compile(r"telefone|celular|whatsapp"), lambda c, k: c.telefone),
    (re.compile(r"\bcep\b"), lambda c, k: c.endereco.cep),
    (
        re.compile(r"endereco|logradouro|\brua\b|bairro|residen|domicili"),
        lambda c, k: _rua(c.endereco),
    ),
    (
        re.compile(r"nome da cidade|forum|comarca"),
        lambda c, k: c.endereco.cidade,
    ),  # fórum: só a cidade
    (re.compile(r"cidade|municipio"), lambda c, k: f"{c.endereco.cidade}/{c.endereco.uf}"),
    (re.compile(r"\bmarca\b"), lambda c, k: c.veiculo_marca),
    (re.compile(r"modelo"), lambda c, k: c.veiculo_modelo),
    (re.compile(r"\bplaca\b"), lambda c, k: c.placa),
    (re.compile(r"\bcor\b"), lambda c, k: c.veiculo_cor),
    (re.compile(r"\bano\b|ano/modelo|ano do"), lambda c, k: c.veiculo_ano),
    (re.compile(r"benefici"), lambda c, k: c.beneficio),
    (re.compile(r"contrato"), lambda c, k: c.contrato),
    (re.compile(r"agencia|conta|deposito"), lambda c, k: c.agencia_conta),
    (re.compile(r"quantidade|parcelas"), lambda c, k: str(c.parcelas)),
    (re.compile(r"data"), lambda c, k: c.data),
    (
        re.compile(
            r"valor|total|gasto|parcela|credito|divida|troco|importancia|emprestimo|desconto|dano moral|restituicao|quantia|soma"
        ),
        lambda c, k: c.money(k),
    ),
    (re.compile(r"local.*avaria|avaria"), lambda c, k: c.avarias),
    (re.compile(r"local"), lambda c, k: c.local),
    (re.compile(r"nacionalidade"), lambda c, k: c.current.nacionalidade),
    (re.compile(r"estado civil"), lambda c, k: c.current.estado_civil),
    (re.compile(r"profiss"), lambda c, k: c.current.profissao),
    (re.compile(r"filia"), lambda c, k: c.current.filiacao),
    (
        re.compile(
            r"resumo|o que aconteceu|fato|narr|relato|descri|pedido|forma de pagamento|meios|contato|causas"
        ),
        lambda c, k: c.fatos or c.outras_info,
    ),
    (re.compile(r"informac"), lambda c, k: c.outras_info),
]


def _ascii_lower(text: str) -> str:
    nfkd = unicodedata.normalize("NFKD", text)
    return "".join(ch for ch in nfkd if not unicodedata.combining(ch)).lower()


def _semantic_key(token: str) -> str:
    """Normaliza o texto do slot a uma chave semântica estável.

    Remove delimitadores, o verbo "digite/digitar [aqui]" e tudo após o primeiro
    exemplo/parêntese/vírgula — de modo que '<digitar o valor total do dano>' e
    '<digite o valor total do dano>' colapsem na MESMA chave (mesmo valor).
    """
    key = _ascii_lower(token).strip("<> .")
    key = re.sub(r"^digit(e|ar)\s+(aqui\s+)?(o|a|os|as)?\s*", "", key)
    key = re.split(r"[,(]|\bex\.?:|\bexemplo|\btais como|\bpor exemplo", key)[0]
    return key.strip()


def _resolve(token: str, ctx: _FillCtx) -> str:
    """Mapeia o texto de um slot ao valor sintético.

    Sem cache: cada gerador é determinístico dada a chave e o estado de `ctx` —
    e os dados pessoais dependem da *parte corrente*, que muda ao longo da peça,
    então não podem ser cacheados por chave.
    """
    key = _semantic_key(token)
    value = next((gen(ctx, key) for pat, gen in _RULES if pat.search(key)), None)
    if value is None:
        value = ctx.autor_nome if "nome" in key else "(não informado)"
    return value


def _num_for(paragraph_low: str, ctx: _FillCtx) -> str:
    """Número de documento adequado ao contexto de um blank 'nº   ' sem rótulo."""
    if "cnpj" in paragraph_low:
        return ctx.reu_cnpj
    if "cpf" in paragraph_low:
        return ctx.current.cpf
    if "benefic" in paragraph_low:
        return ctx.beneficio
    return ctx.contrato


# --- Manipulação de docx (write-back preservando formatação) ---------------------


def _iter_paragraphs(doc: DocxDocument) -> Iterator[Paragraph]:
    """Todos os parágrafos do corpo E das tabelas (recursivo)."""
    yield from doc.paragraphs
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterator[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _replace_in_paragraph(
    paragraph: Paragraph, pattern: re.Pattern[str], repl: Callable[[re.Match[str]], str]
) -> int:
    """Substitui as ocorrências de `pattern` no parágrafo, **cruzando runs**.

    Preserva a formatação da run onde cada ocorrência começa (escreve o valor ali e
    esvazia o trecho casado nas runs seguintes). Processa da direita p/ a esquerda
    para manter os offsets válidos. Retorna o número de substituições.
    """
    runs = paragraph.runs
    if not runs:
        return 0
    full = "".join(r.text for r in runs)
    matches = list(pattern.finditer(full))
    if not matches:
        return 0
    bounds: list[tuple[int, int]] = []
    pos = 0
    for r in runs:
        bounds.append((pos, pos + len(r.text)))
        pos += len(r.text)
    for m in reversed(matches):
        s, e = m.start(), m.end()
        replacement = repl(m)
        for run, (rs, re_) in zip(runs, bounds, strict=True):
            if re_ <= s or rs >= e:  # run sem sobreposição com o match
                continue
            local_s = max(s, rs) - rs
            local_e = min(e, re_) - rs
            if rs <= s < re_:  # run que contém o INÍCIO do match → recebe o valor
                run.text = run.text[:local_s] + replacement + run.text[local_e:]
            else:  # runs do meio/fim → só removem o trecho casado
                run.text = run.text[:local_s] + run.text[local_e:]
    return len(matches)


def _delete_paragraph(paragraph: Paragraph) -> None:
    """Remove o parágrafo do documento — com segurança OOXML.

    Toda célula de tabela (`<w:tc>`) e o corpo precisam de ao menos um `<w:p>`, e o
    parágrafo que carrega o `sectPr` não pode ser removido (corromperia o arquivo —
    o Word recusa abrir). Nesses casos, esvazia o conteúdo em vez de remover o nó.
    """
    element = paragraph._element
    parent = element.getparent()
    has_sectpr = (
        element.find(qn("w:pPr")) is not None
        and element.find(f"{qn('w:pPr')}/{qn('w:sectPr')}") is not None
    )
    siblings = parent.findall(qn("w:p"))
    if len(siblings) > 1 and not has_sectpr:
        parent.remove(element)
    else:
        paragraph.clear()  # mantém o <w:p> (e suas propriedades), zera o conteúdo


def _is_meta(text: str) -> bool:
    low = _ascii_lower(text.strip())
    if not low:
        return False
    if low.startswith("*"):
        return True
    return any(marker in low for marker in _META_MARKERS)


# --- Normalização de cor (evita branco-no-branco do template) --------------------


def _fill_luminance(fill: str | None) -> float:
    """Luminância [0,1] de um fill hex 'RRGGBB'. Fundo branco/ausente → 1.0."""
    if not fill or fill in ("auto", "FFFFFF"):
        return 1.0
    try:
        r, g, b = int(fill[0:2], 16), int(fill[2:4], 16), int(fill[4:6], 16)
    except ValueError:
        return 1.0
    return (0.299 * r + 0.587 * g + 0.114 * b) / 255


def _cell_is_dark(cell_tc) -> bool:
    tc_pr = cell_tc.tcPr
    if tc_pr is None:
        return False
    shd = tc_pr.find(qn("w:shd"))
    fill = shd.get(qn("w:fill")) if shd is not None else None
    return _fill_luminance(fill) < 0.5


def _recolor(paragraph: Paragraph, *, white: bool) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = _WHITE if white else _BLACK


def _normalize_colors(doc: DocxDocument) -> None:
    """Força texto preto (branco só sobre célula escura).

    Os templates do TJDFT trazem títulos de seção em fonte branca herdada de
    estilo, sem faixa de fundo — o que renderiza branco-no-branco (invisível e
    falso-positivo do firewall). Aqui isso é neutralizado para gerar peças limpas.
    """
    for paragraph in doc.paragraphs:
        _recolor(paragraph, white=False)
    for table in doc.tables:
        _normalize_table_colors(table)


def _normalize_table_colors(table: Table) -> None:
    for row in table.rows:
        for cell in row.cells:
            dark = _cell_is_dark(cell._tc)
            for paragraph in cell.paragraphs:
                _recolor(paragraph, white=dark)
            for nested in cell.tables:
                _normalize_table_colors(nested)


# --- Preenchimento de um documento -----------------------------------------------


def _fill_document(doc: DocxDocument, ctx: _FillCtx, *, normalize_colors: bool = True) -> None:
    """Aplica todos os slots in-place no documento."""
    paragraphs = list(_iter_paragraphs(doc))

    # 1. Remove meta-instruções e resolve blocos alternativos ("Situação 1/2").
    alt_groups: dict[str, list[Paragraph]] = {}
    for p in paragraphs:
        text = p.text.strip()
        if _is_meta(text):
            _delete_paragraph(p)
            continue
        alt = re.match(
            r"(situa[çc][ãa]o|op[çc][ãa]o|hip[óo]tese)\s*\d+\s*[-:\u2013]", _ascii_lower(text)
        )
        if alt:
            alt_groups.setdefault("default", []).append(p)
    for group in alt_groups.values():
        keep = ctx.rng.randrange(len(group))
        for i, p in enumerate(group):
            if i != keep:
                _delete_paragraph(p)

    # 2. Preenche placeholders, campos em branco, horário e marca checkboxes.
    for p in list(_iter_paragraphs(doc)):  # relê (parágrafos podem ter sido removidos)
        low = _ascii_lower(p.text)
        _replace_in_paragraph(p, _ANGLE, lambda m: _resolve(m.group(0), ctx))
        _replace_in_paragraph(
            p, _HORA, lambda m: f" {ctx.rng.randint(7, 20):02d}h{ctx.rng.randint(0, 59):02d}min"
        )
        # Sem vírgula no fim: o template já traz a pontuação após o campo em branco.
        _replace_in_paragraph(p, _BLANK, lambda m: f"{m.group(1)}: {_resolve(m.group(1), ctx)}")
        _replace_in_paragraph(p, _RS_BLANK, lambda m, t=low: f"R$ {ctx.money('rs_' + t[:24])} ")
        _replace_in_paragraph(p, _NUM_BLANK, lambda m, t=low: f"nº {_num_for(t, ctx)} ")

    # 2b. Placeholders que cruzam fronteira de parágrafo (abrem '<' num parágrafo e
    #     fecham '>' no seguinte) — a varredura por parágrafo acima não os alcança.
    _fill_cross_paragraph(doc, ctx)

    _select_checkboxes(doc, ctx)

    # 3. Neutraliza branco-no-branco herdado do template (após preencher/marcar).
    if normalize_colors:
        _normalize_colors(doc)


def _fill_cross_paragraph(doc: DocxDocument, ctx: _FillCtx) -> None:
    """Resolve placeholders `<...>` partidos entre parágrafos consecutivos.

    Junta o texto do parágrafo que abre '<' (sem fechar) com os seguintes até achar
    '>', resolve o placeholder inteiro, escreve o valor no parágrafo de abertura e
    remove o trecho consumido dos demais (lookahead de até 3 parágrafos).
    """
    paragraphs = list(_iter_paragraphs(doc))
    for i, opener in enumerate(paragraphs):
        text = opener.text
        lt, gt = text.rfind("<"), text.rfind(">")
        if lt == -1 or lt < gt:  # nenhum '<' aberto e pendente
            continue
        # Escaneia (SEM mutar) os próximos parágrafos atrás do '>' de fechamento.
        # Um novo '<' antes do '>' significa que ESTE placeholder é malformado (sem
        # fechamento) e o '>' à frente é de outro — não pareia com ele.
        combined = text[lt:]
        close_j = None
        for j in range(i + 1, min(i + 4, len(paragraphs))):
            nxt = paragraphs[j].text
            k = nxt.find(">")
            if "<" in (nxt if k == -1 else nxt[:k]):
                break  # novo placeholder abriu antes do fechamento → malformado
            if k != -1:
                combined += nxt[: k + 1]
                close_j = j
                break
            combined += nxt
        if close_j is None:  # malformado / sem '>' no lookahead: resolve só o aberto
            value = _resolve(text[lt:], ctx)
            _replace_in_paragraph(opener, _OPEN_ANGLE, lambda m, v=value: v)
            continue
        value = _resolve(combined, ctx)
        _replace_in_paragraph(opener, _OPEN_ANGLE, lambda m, v=value: v)
        for m in range(i + 1, close_j):
            _replace_in_paragraph(paragraphs[m], re.compile(r".+"), lambda mm: "")
        _replace_in_paragraph(paragraphs[close_j], _CLOSE_ANGLE, lambda mm: "")


def _mark_one(paragraph: Paragraph, occurrence: int) -> None:
    """Marca apenas a `occurrence`-ésima checkbox (0-based) do parágrafo."""
    full = "".join(r.text for r in paragraph.runs)
    starts = [m.start() for m in _CHECKBOX.finditer(full)]
    if occurrence >= len(starts):
        return
    target = starts[occurrence]
    _replace_in_paragraph(
        paragraph, _CHECKBOX, lambda m: "( X )" if m.start() == target else m.group(0)
    )


def _select_checkboxes(doc: DocxDocument, ctx: _FillCtx) -> None:
    """Marca **exatamente uma** opção por grupo de checkboxes.

    Um grupo é uma corrida contígua de parágrafos com checkbox; opções inline
    ("(  ) X ou (  ) Y" no mesmo parágrafo) também contam como opções do grupo.
    """
    group: list[tuple[Paragraph, int]] = []  # (parágrafo, nº de checkboxes nele)

    def flush() -> None:
        if not group:
            return
        total = sum(n for _, n in group)
        pick = ctx.rng.randrange(total)
        acc = 0
        for paragraph, n in group:
            if pick < acc + n:
                _mark_one(paragraph, pick - acc)
                break
            acc += n
        group.clear()

    for p in _iter_paragraphs(doc):
        boxes = len(_CHECKBOX.findall(p.text))
        if boxes:
            group.append((p, boxes))
        else:
            flush()
    flush()


# --- Render: docx -> pdf via LibreOffice -----------------------------------------


def _docx_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError(
            "LibreOffice (soffice) não encontrado. Instale-o para gerar o PDF:\n"
            "    sudo apt-get install -y libreoffice-writer"
        )
    with tempfile.TemporaryDirectory() as profile:
        subprocess.run(
            [
                soffice,
                "--headless",
                f"-env:UserInstallation=file://{profile}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            timeout=120,
        )
    pdf = out_dir / f"{docx_path.stem}.pdf"
    if not pdf.exists():
        raise RuntimeError(f"Conversão para PDF não produziu {pdf}")
    return pdf


# --- Ground truth ----------------------------------------------------------------


@dataclass
class _Label:
    category: str
    description: str
    is_malicious: bool = False
    expected_verdict: str = "PASS"
    vector: str | None = None
    rito: str = "CIVEL"
    expect_liminar: bool | None = None
    expect_semaforo: str | None = None
    expect_requer_emenda: bool | None = None


def _slugify(name: str) -> str:
    base = _ascii_lower(Path(name).stem)
    base = re.sub(r"^[\d.\s]+", "", base)  # remove numeração ("2.1 ")
    base = re.sub(r"[^a-z0-9]+", "_", base).strip("_")
    return base[:60] or "peticao"


# --- LLM (enriquecimento opcional da narrativa) ----------------------------------


def _enrich_with_llm(ctx: _FillCtx, orientacao: str | None) -> None:
    """Preenche `ctx.fatos`/`ctx.outras_info` com narrativa coerente via LLMProvider.

    Lazy: só importa config/factory quando `--enrich-llm` é usado.
    """
    import asyncio

    from pydantic import BaseModel, Field

    from sherpi.config import Settings
    from sherpi.infrastructure.llm.factory import build_llm_provider
    from sherpi.shared_kernel.ports import ChatMessage

    class _Narrativa(BaseModel):
        fatos: str = Field(
            description="Narrativa dos fatos, 1-2 parágrafos, tom forense, em pt-BR."
        )
        outras_informacoes: str = Field(description="Frase curta com informações pós-fato.")

    settings = Settings()
    provider = build_llm_provider(settings)
    contexto = f"\n\nOrientações do tribunal:\n{orientacao[:4000]}" if orientacao else ""
    prompt = (
        "Você redige a seção DOS FATOS de uma petição inicial SINTÉTICA (dados fictícios, "
        f"sem PII real). Autor: {ctx.autor_nome}. Réu: {ctx.reu_nome}. "
        f"Valor do dano: R$ {_money_brl(ctx.valor_base)}. Local: {ctx.local}. Avarias: {ctx.avarias}."
        f"{contexto}"
    )
    result = asyncio.run(provider.complete([ChatMessage(role="user", content=prompt)], _Narrativa))
    ctx.fatos = result.fatos
    ctx.outras_info = result.outras_informacoes


# --- Orquestração / CLI ----------------------------------------------------------


def generate(
    template: Path,
    *,
    out_dir: Path,
    n: int = 1,
    seed: int = 42,
    rito: str = "CIVEL",
    category: str = "clean",
    orientacao: Path | None = None,
    enrich_llm: bool = False,
    keep_docx: bool = True,
) -> dict[str, dict]:
    """Gera `n` petições a partir de `template`. Retorna o mapa de labels gerado."""
    out_dir.mkdir(parents=True, exist_ok=True)
    slug = _slugify(template.name)
    orientacao_text = (
        _read_pdf_text(orientacao) if (enrich_llm and orientacao and orientacao.exists()) else None
    )
    labels: dict[str, dict] = {}

    for i in range(n):
        ctx = _FillCtx(rng=random.Random(seed + i))
        if enrich_llm:
            _enrich_with_llm(ctx, orientacao_text)
        doc = docx.Document(str(template))
        _fill_document(doc, ctx)

        stem = f"{category}_{slug}_v{i + 1}"
        docx_path = out_dir / f"{stem}.docx"
        doc.save(str(docx_path))
        _docx_to_pdf(docx_path, out_dir)
        if not keep_docx:
            docx_path.unlink()

        labels[f"{stem}.pdf"] = asdict(
            _Label(
                category=category,
                description=f"Petição sintética de template: {template.name}",
                rito=rito,
            )
        )
        print(f"  ✓ {stem}.pdf" + (f" (+ {stem}.docx)" if keep_docx else ""))

    return labels


def _read_pdf_text(path: Path) -> str:
    import pymupdf

    with pymupdf.open(str(path)) as doc:
        return "\n".join(page.get_text() for page in doc)


def _merge_labels(out_dir: Path, new: dict[str, dict]) -> None:
    """Mescla os novos labels no labels.json existente (preserva o corpus)."""
    path = out_dir / "labels.json"
    existing = json.loads(path.read_text(encoding="utf-8")) if path.exists() else {}
    existing.update(new)
    path.write_text(json.dumps(existing, indent=2, ensure_ascii=False), encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Gera petições sintéticas a partir de um template .docx."
    )
    parser.add_argument(
        "template", type=Path, help="template .docx (ex.: docs/templates/2.1 ...docx)"
    )
    parser.add_argument(
        "--orientacao",
        type=Path,
        default=None,
        help="PDF de orientações (contexto p/ --enrich-llm)",
    )
    parser.add_argument("--n", type=int, default=1, help="quantas petições gerar")
    parser.add_argument("--seed", type=int, default=42, help="semente base do RNG")
    parser.add_argument(
        "--out", type=Path, default=Path("data/synthetic"), help="diretório de saída"
    )
    parser.add_argument("--rito", default="CIVEL", help="rito (CIVEL | TRABALHISTA) para o label")
    parser.add_argument(
        "--category", default="clean", help="categoria do label (clean | defect | ...)"
    )
    parser.add_argument(
        "--enrich-llm", action="store_true", help="reescreve a narrativa via LLMProvider"
    )
    parser.add_argument("--no-docx", action="store_true", help="descarta o .docx, mantém só o .pdf")
    args = parser.parse_args()

    print(f"Template: {args.template.name}")
    labels = generate(
        args.template,
        out_dir=args.out,
        n=args.n,
        seed=args.seed,
        rito=args.rito,
        category=args.category,
        orientacao=args.orientacao,
        enrich_llm=args.enrich_llm,
        keep_docx=not args.no_docx,
    )
    _merge_labels(args.out, labels)
    print(f"{len(labels)} petição(ões) + labels.json atualizados em {args.out}/")


if __name__ == "__main__":
    main()
