"""Testes do gerador de petições a partir de templates (`synthetic.from_template`).

Usam um mini-template construído em memória (não dependem do LibreOffice nem dos
arquivos em docs/templates) e exercitam a engine determinística: detecção de slots,
write-back cruzando runs, consistência de valores e marcação de checkbox.
"""

from __future__ import annotations

import random
import re

import docx
from synthetic.from_template import (
    _ANGLE,
    _fill_document,
    _FillCtx,
    _iter_paragraphs,
    _replace_in_paragraph,
    _semantic_key,
)


def _mini_template() -> docx.document.Document:
    """Template sintético com os 4 tipos de slot + texto branco a normalizar."""
    doc = docx.Document()
    doc.add_paragraph("AO JUIZADO ESPECIAL CÍVEL DE <DIGITE O NOME DA CIDADE> - DF.")
    doc.add_paragraph(
        "REQUERENTE: <DIGITE O NOME DO AUTOR>, nacionalidade:      , estado civil:      ,"
    )
    doc.add_paragraph("em face de <DIGITE O NOME DO RÉU>.")
    doc.add_paragraph("Atribui à causa o valor de R$ <digitar o valor total do dano>.")
    doc.add_paragraph("Condenar ao pagamento de R$ <digite o valor total do dano>.")
    doc.add_paragraph("(  ) opção A")
    doc.add_paragraph("(  ) opção B")
    doc.add_paragraph("(sugestões de texto - adote um e apague o outro):")  # meta → removido
    return doc


def _all_text(doc: docx.document.Document) -> str:
    return "\n".join(p.text for p in _iter_paragraphs(doc))


def test_fill_preenche_todos_os_slots() -> None:
    doc = _mini_template()
    _fill_document(doc, _FillCtx(rng=random.Random(7)))
    text = _all_text(doc)

    assert "<" not in text and ">" not in text  # nenhum placeholder remanescente
    assert "(não informado)" not in text  # nenhum fallback
    assert "( X )" in text  # uma opção marcada
    assert text.count("( X )") == 1  # exatamente uma por grupo
    assert "sugestões de texto" not in text  # meta-instrução removida


def test_valores_consistentes_e_partes_distintas() -> None:
    doc = _mini_template()
    ctx = _FillCtx(rng=random.Random(7))
    _fill_document(doc, ctx)
    text = _all_text(doc)

    # 'valor total do dano' aparece 2x (digite/digitar) → MESMO valor.
    dano = ctx.money("valor total do dano")
    assert text.count(dano) == 2
    # autor e réu são pessoas distintas e ambos presentes.
    assert ctx.autor_nome != ctx.reu_nome
    assert ctx.autor_nome in text and ctx.reu_nome in text
    # campos em branco rotulados preenchidos (estado civil vem do catálogo → varia).
    assert "nacionalidade: brasileiro(a)" in text
    assert re.search(r"estado civil: \S", text)


def test_determinismo_por_seed() -> None:
    a = _mini_template()
    b = _mini_template()
    _fill_document(a, _FillCtx(rng=random.Random(42)))
    _fill_document(b, _FillCtx(rng=random.Random(42)))
    assert _all_text(a) == _all_text(b)


def test_semantic_key_colapsa_digite_digitar() -> None:
    assert _semantic_key("<digitar o valor total do dano>") == _semantic_key(
        "<digite o valor total do dano>"
    )
    assert _semantic_key("<digite aqui o nº benefício INSS>") == "no beneficio inss"


def test_replace_cruza_runs_preservando_demais() -> None:
    doc = docx.Document()
    p = doc.add_paragraph()
    # placeholder partido em 3 runs, como ocorre nos templates reais.
    p.add_run("Nome: <DIGITE ")
    p.add_run("O NOME")
    p.add_run("> fim")
    n = _replace_in_paragraph(p, _ANGLE, lambda m: "Fulano")
    assert n == 1
    assert p.text == "Nome: Fulano fim"


def test_money_estavel_por_chave() -> None:
    ctx = _FillCtx(rng=random.Random(3))
    assert ctx.money("parcela") == ctx.money("parcela")  # estável
    # chaves distintas tendem a divergir (âncora * fatores diferentes).
    distintos = {ctx.money(k) for k in ("parcela", "total", "dano moral", "troco", "credito")}
    assert len(distintos) >= 2


def test_litisconsorcio_partes_distintas() -> None:
    """Em litisconsórcio (proprietário/condutor), cada parte tem nome e CPF próprios."""
    doc = docx.Document()
    doc.add_paragraph(
        "AUTOR-PROPRIETÁRIO: <digite o nome do autor-proprietário>, "
        "inscrito no CPF sob o nº:      ,"
    )
    doc.add_paragraph(
        "AUTOR-CONDUTOR: <digite o nome do autor-condutor>, inscrito no CPF sob o nº:      ,"
    )
    ctx = _FillCtx(rng=random.Random(9))
    _fill_document(doc, ctx)
    text = _all_text(doc)

    assert ctx.personas[0].nome in text and ctx.personas[1].nome in text
    assert ctx.personas[0].nome != ctx.personas[1].nome
    cpfs = re.findall(r"\d{3}\.\d{3}\.\d{3}-\d{2}", text)
    assert len(cpfs) == 2 and cpfs[0] != cpfs[1]  # CPFs distintos por parte


def test_remocao_em_celula_unica_nao_esvazia(tmp_path) -> None:
    """Remover a única linha (meta) de uma célula NÃO deixa o <w:tc> sem <w:p>
    (OOXML inválido → Word recusa abrir)."""
    doc = docx.Document()
    cell = doc.add_table(rows=1, cols=1).cell(0, 0)
    cell.paragraphs[0].add_run("(sugestões de texto - apague o outro):")  # meta → seria removido
    _fill_document(doc, _FillCtx(rng=random.Random(1)))

    out = tmp_path / "out.docx"
    doc.save(str(out))
    reopened = docx.Document(str(out))  # reabrir valida o OOXML
    assert reopened.tables[0].cell(0, 0).paragraphs  # célula mantém ao menos um <w:p>


def test_diferencia_pessoa_fisica_de_juridica() -> None:
    """Réu 'requerida' genérico = pessoa física (nome+CPF); só vira empresa com
    indicador explícito de PJ (razão social/empresa)."""
    fisica = docx.Document()
    fisica.add_paragraph("em face de <DIGITE O NOME DA PARTE REQUERIDA>.")
    juridica = docx.Document()
    juridica.add_paragraph("em face de <DIGITE A RAZÃO SOCIAL DA EMPRESA REQUERIDA>.")

    cf, cj = _FillCtx(rng=random.Random(7)), _FillCtx(rng=random.Random(7))
    _fill_document(fisica, cf)
    _fill_document(juridica, cj)

    assert cf.reu_nome in _all_text(fisica)  # pessoa física: nome de pessoa
    assert cj.reu_empresa in _all_text(juridica)  # PJ: razão social da empresa


def test_placeholder_cruza_paragrafo() -> None:
    """Placeholder que abre '<' num parágrafo e fecha '>' no seguinte é resolvido."""
    doc = docx.Document()
    doc.add_paragraph('AÇÃO DE <digite o nome da ação ou "CONHECIMENTO"')
    doc.add_paragraph("> proposta em face de quem de direito.")
    _fill_document(doc, _FillCtx(rng=random.Random(2)))
    text = _all_text(doc)
    assert "<" not in text and "conhecimento" in text.lower()
